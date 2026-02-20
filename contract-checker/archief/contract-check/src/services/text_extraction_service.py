"""
Text Extraction Service - Extract plain text from contract files
"""
from io import BytesIO
from typing import Optional
from pathlib import Path

# Document processing imports
from docx import Document  # python-docx
from openpyxl import load_workbook  # openpyxl
try:
    import pdfplumber  # For PDF extraction
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False
    # Fallback to PyPDF2 if pdfplumber not available
    try:
        from PyPDF2 import PdfReader
        HAS_PYPDF2 = True
    except ImportError:
        HAS_PYPDF2 = False


class TextExtractionError(Exception):
    """Custom exception for text extraction errors"""
    pass


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file.

    Args:
        file_content: Binary PDF content

    Returns:
        Extracted plain text

    Raises:
        TextExtractionError: If extraction fails
    """
    if not (HAS_PDFPLUMBER or HAS_PYPDF2):
        raise TextExtractionError(
            "Geen PDF library beschikbaar. Installeer pdfplumber of PyPDF2."
        )

    try:
        pdf_file = BytesIO(file_content)

        if HAS_PDFPLUMBER:
            # Preferred: pdfplumber (better text extraction)
            with pdfplumber.open(pdf_file) as pdf:
                text_parts = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                return '\n\n'.join(text_parts)
        else:
            # Fallback: PyPDF2
            reader = PdfReader(pdf_file)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return '\n\n'.join(text_parts)

    except Exception as e:
        raise TextExtractionError(f"PDF extractie fout: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX file.

    Args:
        file_content: Binary DOCX content

    Returns:
        Extracted plain text

    Raises:
        TextExtractionError: If extraction fails
    """
    try:
        docx_file = BytesIO(file_content)
        doc = Document(docx_file)

        # Extract text from paragraphs
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))

        return '\n\n'.join(text_parts)

    except Exception as e:
        raise TextExtractionError(f"DOCX extractie fout: {str(e)}")


def extract_text_from_xlsx(file_content: bytes) -> str:
    """
    Extract text from XLSX file.

    Args:
        file_content: Binary XLSX content

    Returns:
        Extracted plain text (sheet by sheet)

    Raises:
        TextExtractionError: If extraction fails
    """
    try:
        xlsx_file = BytesIO(file_content)
        workbook = load_workbook(xlsx_file, data_only=True)

        text_parts = []

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]

            # Add sheet name as header
            text_parts.append(f"=== {sheet_name} ===")

            # Extract all cell values
            for row in sheet.iter_rows(values_only=True):
                # Filter out empty cells and convert to strings
                row_values = [str(cell) for cell in row if cell is not None and str(cell).strip()]
                if row_values:
                    text_parts.append(' | '.join(row_values))

        return '\n\n'.join(text_parts)

    except Exception as e:
        raise TextExtractionError(f"XLSX extractie fout: {str(e)}")


def extract_text(file_content: bytes, filename: str) -> str:
    """
    Extract text from file based on extension.

    Args:
        file_content: Binary file content
        filename: Original filename (to determine type)

    Returns:
        Extracted plain text

    Raises:
        TextExtractionError: If extraction fails or unsupported file type
    """
    ext = Path(filename).suffix.lower()

    extractors = {
        '.pdf': extract_text_from_pdf,
        '.docx': extract_text_from_docx,
        '.xlsx': extract_text_from_xlsx,
    }

    extractor = extractors.get(ext)
    if not extractor:
        raise TextExtractionError(
            f"Niet ondersteund bestandstype: {ext}. "
            f"Ondersteund: {', '.join(extractors.keys())}"
        )

    return extractor(file_content)


def save_extracted_text(db, file_id: int, extracted_text: str) -> None:
    """
    Save extracted text to database.

    Args:
        db: Database session
        file_id: File ID
        extracted_text: Extracted plain text
    """
    from sqlalchemy import text

    db.execute(
        text("""
            UPDATE contract_checker.contract_files
            SET extracted_text = :extracted_text,
                updated_at = NOW()
            WHERE id = :file_id
        """),
        {
            'file_id': file_id,
            'extracted_text': extracted_text
        }
    )
    db.commit()


def get_extracted_text(db, file_id: int) -> Optional[str]:
    """
    Get extracted text from database.

    Args:
        db: Database session
        file_id: File ID

    Returns:
        Extracted text or None if not found
    """
    from sqlalchemy import text

    result = db.execute(
        text("""
            SELECT extracted_text FROM contract_checker.contract_files
            WHERE id = :file_id AND active = true
        """),
        {'file_id': file_id}
    )
    row = result.fetchone()
    return row[0] if row else None
