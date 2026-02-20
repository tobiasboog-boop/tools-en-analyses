from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document  # type: ignore[import-untyped]
import openpyxl
from src.config import config
from src.models import db, Contract
from src.models.contract_relatie import ContractRelatie


class ContractLoader:
    """Service to load and parse contract documents.

    Integrates with the contracts metadata table to look up
    which contract file belongs to which client.

    Supports two modes:
    1. Pre-extracted: Uses .txt files from the 'extracted' subfolder (preferred)
    2. On-the-fly: Extracts from .docx/.xlsx if no pre-extracted file exists
    """

    def __init__(self, folder_path: str = None, prefer_extracted: bool = True):
        self.folder_path = Path(folder_path or config.CONTRACTS_FOLDER)
        self.extracted_folder = self.folder_path / "llm_conversie"
        self.prefer_extracted = prefer_extracted
        self._content_cache: Dict[str, str] = {}  # filename -> content

    def get_contract_for_debiteur(self, debiteur_code: str) -> Optional[Dict[str, Any]]:
        """Get contract info and content for a specific debiteur.

        Looks up the contract via contract_relatie table, then loads the file content.
        Returns None if no active contract found for this debiteur.
        """
        from sqlalchemy import text

        session = db()
        try:
            # Find contract via contract_relatie table (client_id = debiteur_code)
            result = session.execute(text("""
                SELECT c.id, c.filename, c.source_file, c.source_sheet, c.llm_ready, c.content
                FROM contract_checker.contracts c
                JOIN contract_checker.contract_relatie cr ON c.id = cr.contract_id
                WHERE cr.client_id = :debiteur_code
                  AND c.active = true
                LIMIT 1
            """), {"debiteur_code": debiteur_code})

            row = result.fetchone()
            if not row:
                return None

            contract_id, filename, source_file, source_sheet, llm_ready, content = row

            # Use llm_ready if available, otherwise content
            text_content = llm_ready if llm_ready else content

            if text_content is None:
                return None

            return {
                "id": contract_id,
                "filename": filename,
                "source_file": source_file,
                "source_sheet": source_sheet,
                "content": text_content,
                "llm_ready": bool(llm_ready),
            }
        finally:
            session.close()

    def get_contract_for_debiteur(self, debiteur_code: str) -> Optional[Dict[str, Any]]:
        """Get contract info for a debiteur via contract_relatie koppeling.

        Looks up the contract_relatie for the debiteur_code, then loads the contract.
        Returns None if no active contract found for this debiteur.
        """
        session = db()
        try:
            # Find the contract_relatie for this debiteur
            relatie = session.query(ContractRelatie).filter(
                ContractRelatie.client_id == debiteur_code
            ).first()

            if not relatie:
                return None

            # Get the contract by id
            contract = session.query(Contract).filter(
                Contract.id == relatie.contract_id,
                Contract.active.is_(True)
            ).first()

            if not contract:
                return None

            # Use content from database (llm_ready or content)
            content = contract.llm_ready or contract.content

            return {
                "id": contract.id,
                "filename": contract.filename,
                "client_id": relatie.client_id,
                "client_name": relatie.client_name,
                "content": content,
            }
        finally:
            session.close()

    def _load_contract_content(self, filename: str) -> Optional[str]:
        """Load contract file content, with caching.

        Priority:
        1. If prefer_extracted=True, look for pre-extracted .txt file first
        2. Fall back to on-the-fly extraction from .docx/.xlsx
        """
        # Check cache first (useful for group contracts)
        if filename in self._content_cache:
            return self._content_cache[filename]

        content = None
        source = None

        # Try pre-extracted .txt file first
        if self.prefer_extracted:
            extracted_path = self._get_extracted_path(filename)
            if extracted_path and extracted_path.exists():
                try:
                    content = extracted_path.read_text(encoding="utf-8")
                    source = "extracted"
                except Exception as e:
                    print(f"Warning: Could not read extracted file {extracted_path}: {e}")

        # Fall back to on-the-fly extraction
        if content is None:
            filepath = self.folder_path / filename
            if not filepath.exists():
                print(f"Warning: Contract file not found: {filepath}")
                return None

            try:
                ext = filepath.suffix.lower()
                if ext == ".docx":
                    content = self._load_docx(str(filepath))
                    source = "docx"
                elif ext == ".xlsx":
                    content = self._load_xlsx(str(filepath))
                    source = "xlsx"
                else:
                    print(f"Warning: Unsupported file type: {ext}")
                    return None
            except Exception as e:
                print(f"Error loading {filename}: {e}")
                return None

        if content:
            # Cache the content
            self._content_cache[filename] = content
            if source == "extracted":
                print(f"  Using pre-extracted: {filename}")

        return content

    def _get_extracted_path(self, filename: str) -> Optional[Path]:
        """Get path to pre-extracted .txt file for a contract."""
        if not self.extracted_folder.exists():
            return None
        # Convert original filename to .txt
        stem = Path(filename).stem
        return self.extracted_folder / f"{stem}.txt"

    def load_all_contracts(self) -> List[Dict[str, Any]]:
        """Load all active contracts from database.

        Returns list of contracts with their content.
        Uses caching so each file is only read once (for group contracts).
        """
        session = db()
        try:
            contracts = session.query(Contract).filter(
                Contract.active.is_(True)
            ).all()

            results = []
            for contract in contracts:
                content = self._load_contract_content(contract.filename)
                if content:
                    results.append({
                        "id": contract.id,
                        "filename": contract.filename,
                        "source_file": contract.source_file,
                        "source_sheet": contract.source_sheet,
                        "content": content,
                        "llm_ready": contract.llm_ready,
                    })
            return results
        finally:
            session.close()

    def get_unique_contract_files(self) -> List[Dict[str, Any]]:
        """Get unique contract files with their content.

        For group contracts, returns each file only once.
        Useful for sending to Claude (avoid duplicating content).
        Uses llm_ready version if available, otherwise falls back to content.
        """
        session = db()
        try:
            # Get active contracts
            contracts = session.query(Contract).filter(
                Contract.active.is_(True)
            ).all()

            results = []
            for contract in contracts:
                # Prefer llm_ready over content
                text_content = contract.llm_ready if contract.llm_ready else contract.content

                if text_content:
                    results.append({
                        "filename": contract.filename,
                        "content": text_content,
                        "content_length": len(text_content),
                        "llm_ready": bool(contract.llm_ready),
                    })
            return results
        finally:
            session.close()

    def get_contracts_text(self) -> str:
        """Get all unique contracts as a single text block for Claude.

        Uses llm_ready version when available for better AI performance.
        """
        contracts = self.get_unique_contract_files()

        parts = []
        for contract in contracts:
            parts.append(f"### CONTRACT: {contract['filename']} ###")
            parts.append(contract["content"])
            parts.append("")

        return "\n".join(parts)

    def get_contracts_summary(self) -> List[Dict[str, Any]]:
        """Get summary of unique contract files."""
        return self.get_unique_contract_files()

    def _load_docx(self, path: str) -> str:
        """Extract text content from a Word document."""
        doc = Document(path)
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.replace("|", "").strip():
                    paragraphs.append(row_text)

        return "\n".join(paragraphs)

    def _load_xlsx(self, path: str) -> str:
        """Extract text content from an Excel file."""
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        content_parts = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            content_parts.append(f"=== Sheet: {sheet_name} ===")

            for row in sheet.iter_rows(values_only=True):
                row_values = [
                    str(cell) if cell is not None else ""
                    for cell in row
                ]
                if any(v.strip() for v in row_values):
                    content_parts.append(" | ".join(row_values))

        wb.close()
        return "\n".join(content_parts)

    def clear_cache(self):
        """Clear the content cache."""
        self._content_cache.clear()

    def get_extraction_status(self) -> Dict[str, Any]:
        """Check which contracts have pre-extracted .txt files.

        Returns summary of extraction status for all registered contracts.
        Useful for identifying which contracts still need extraction.
        """
        session = db()
        try:
            # Get all unique filenames from database
            filenames = session.query(Contract.filename).filter(
                Contract.active.is_(True)
            ).distinct().all()

            extracted = []
            not_extracted = []

            for (filename,) in filenames:
                extracted_path = self._get_extracted_path(filename)
                source_path = self.folder_path / filename

                status = {
                    "filename": filename,
                    "source_exists": source_path.exists(),
                }

                if extracted_path and extracted_path.exists():
                    status["extracted_path"] = str(extracted_path)
                    status["extracted_size"] = extracted_path.stat().st_size
                    extracted.append(status)
                else:
                    not_extracted.append(status)

            return {
                "total": len(filenames),
                "extracted_count": len(extracted),
                "not_extracted_count": len(not_extracted),
                "extracted": extracted,
                "not_extracted": not_extracted,
                "extracted_folder": str(self.extracted_folder),
            }
        finally:
            session.close()
