#!/usr/bin/env python3
"""
Extract contracts to LLM-readable text format and generate contract register.

This script:
1. Converts contract files (.docx, .xlsx) to plain text files for LLM processing
2. Generates a contract register template (CSV) for client-contract mapping
3. Optionally: Auto-matches contract filenames to Sintes relaties (fuzzy matching)

Usage:
    # Extract all contracts and generate register template
    python extract_contracts.py

    # With auto-matching to Sintes relaties
    python extract_contracts.py --auto-match

    # Extract specific file(s)
    python extract_contracts.py --files "Contract_A.docx" "Contract_B.xlsx"

    # Only generate register template (no text extraction)
    python extract_contracts.py --register-only

    # Dry run (show what would be done)
    python extract_contracts.py --dry-run

    # Custom output folder
    python extract_contracts.py --output ./extracted_contracts

Output:
    - .txt files in the 'extracted' subfolder (LLM-readable contract text)
    - contract_register_template.csv (template for client-contract mapping)
    - With --auto-match: contract_register_matched.csv (with suggested client_ids)

Workflow:
    1. Notifica runs this script on contracts received from WVC
    2. If --auto-match: Review suggested matches, adjust where needed
    3. Fill in remaining client_id mappings manually
    4. Notifica syncs register to database: python sync_contracts.py register.csv
    5. WVC can review extracted text via the Streamlit UI
"""
import argparse
import csv
import re
import sys
from datetime import datetime
from pathlib import Path

from docx import Document  # type: ignore[import-untyped]
import openpyxl

from src.config import config


# ==============================================================================
# CONTRACT REGISTER GENERATION
# ==============================================================================

def generate_register_template(
    source_folder: Path,
    output_path: Path,
    specific_files: list[str] = None,
    dry_run: bool = False
) -> dict:
    """
    Generate a contract register template CSV.

    Creates a CSV with one row per contract file, with columns:
    - filename: The contract file name (filled in automatically)
    - client_id: To be filled in by Notifica
    - client_name: To be filled in by Notifica
    - contract_number: Optional
    - contract_type: Optional (Standaard, Premium, Basis)
    - start_date: Optional (YYYY-MM-DD)
    - end_date: Optional (YYYY-MM-DD)
    - notes: Optional

    Returns:
        Statistics dict with counts
    """
    stats = {"files_found": 0, "register_created": False}

    # Get files to include
    if specific_files:
        files = [source_folder / f for f in specific_files]
    else:
        files = list(source_folder.glob("*.docx")) + list(source_folder.glob("*.xlsx"))

    # Filter to existing files only
    files = [f for f in files if f.exists()]
    stats["files_found"] = len(files)

    if not files:
        print(f"No contract files found in: {source_folder}")
        return stats

    print(f"Found {len(files)} contract file(s)")
    print(f"Register output: {output_path}")
    print("")

    if dry_run:
        print("[DRY RUN] Would create register with files:")
        for f in sorted(files):
            print(f"  - {f.name}")
        return stats

    # Create register CSV
    fieldnames = [
        "filename",
        "client_id",
        "client_name",
        "contract_number",
        "contract_type",
        "start_date",
        "end_date",
        "notes",
    ]

    with open(output_path, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames, delimiter=";")
        writer.writeheader()

        for filepath in sorted(files):
            writer.writerow({
                "filename": filepath.name,
                "client_id": "",  # To be filled in
                "client_name": "",  # To be filled in
                "contract_number": "",
                "contract_type": "",
                "start_date": "",
                "end_date": "",
                "notes": "",
            })
            print(f"  Added: {filepath.name}")

    stats["register_created"] = True
    print("")
    print(f"Register template created: {output_path}")
    print("")
    print("Next steps:")
    print("  1. Open the CSV and fill in client_id for each contract")
    print("  2. For group contracts, duplicate the row for each client")
    print("  3. Run: python sync_contracts.py <register.csv>")

    return stats


def generate_register_with_matching(
    source_folder: Path,
    output_path: Path,
    specific_files: list[str] = None,
    dry_run: bool = False,
    min_score: float = 0.4,
    top_n: int = 3
) -> dict:
    """
    Generate contract register with auto-matched client suggestions from Sintes.

    Uses fuzzy matching to suggest client_id based on contract filenames.

    Returns:
        Statistics dict with counts and match info
    """
    from src.services.relatie_service import RelatieService, ContractMatcher

    stats = {
        "files_found": 0,
        "register_created": False,
        "relaties_loaded": 0,
        "matched": 0,
        "unmatched": 0
    }

    # Get files to include
    if specific_files:
        files = [source_folder / f for f in specific_files]
    else:
        files = list(source_folder.glob("*.docx")) + list(source_folder.glob("*.xlsx"))

    files = [f for f in files if f.exists()]
    stats["files_found"] = len(files)

    if not files:
        print(f"No contract files found in: {source_folder}")
        return stats

    print(f"Found {len(files)} contract file(s)")
    print(f"Register output: {output_path}")
    print("")

    # Load relaties from Sintes
    print("Loading relaties from Sintes database...")
    service = RelatieService()
    try:
        relaties = service.get_relaties()
        stats["relaties_loaded"] = len(relaties)
        print(f"  Loaded {len(relaties)} relaties")
    except Exception as e:
        print(f"  ERROR loading relaties: {e}")
        print("  Falling back to empty register (no suggestions)")
        relaties = []
    finally:
        service.close()

    if not relaties:
        print("  WARNING: No relaties found. Check database connection.")
        print("  Creating register without suggestions...")

    # Perform matching
    matcher = ContractMatcher(relaties) if relaties else None
    matches = {}

    if matcher:
        print("")
        print("Matching contracts to relaties...")
        for filepath in files:
            contract_matches = matcher.find_matches(
                filepath.name,
                top_n=top_n,
                min_score=min_score
            )
            matches[filepath.name] = contract_matches

            if contract_matches:
                best = contract_matches[0]
                print(f"  {filepath.name}")
                print(f"    -> Best match: {best['client_name']} ({best['client_id']}) "
                      f"[score: {best['score']:.2f}, type: {best['match_type']}]")
                stats["matched"] += 1
            else:
                print(f"  {filepath.name}")
                print(f"    -> No match found")
                stats["unmatched"] += 1

    if dry_run:
        print("")
        print("[DRY RUN] Would create register with matches above")
        return stats

    # Create register CSV with suggestions
    # Note: relatie_key is included for internal linking (to werkbonnen)
    # but client_id (Relatie Code) is what WVC recognizes
    fieldnames = [
        "filename",
        "client_id",
        "client_name",
        "relatie_key",
        "match_score",
        "match_type",
        "matched_on",
        "alternative_1",
        "alternative_2",
        "contract_number",
        "contract_type",
        "start_date",
        "end_date",
        "notes",
    ]

    with open(output_path, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames, delimiter=";")
        writer.writeheader()

        for filepath in sorted(files):
            contract_matches = matches.get(filepath.name, [])

            # Best match (if any)
            best = contract_matches[0] if contract_matches else None

            # Alternative matches (format: "client_id - name (score)")
            alt1 = contract_matches[1] if len(contract_matches) > 1 else None
            alt2 = contract_matches[2] if len(contract_matches) > 2 else None

            def format_alt(m):
                if not m:
                    return ""
                name = m.get('short_name') or m.get('client_name', '')
                return f"{m['client_id']} - {name} ({m['score']:.2f})"

            row = {
                "filename": filepath.name,
                "client_id": best["client_id"] if best else "",
                "client_name": best["client_name"] if best else "",
                "relatie_key": best.get("relatie_key", "") if best else "",
                "match_score": f"{best['score']:.2f}" if best else "",
                "match_type": best["match_type"] if best else "",
                "matched_on": best.get("matched_on", "") if best else "",
                "alternative_1": format_alt(alt1),
                "alternative_2": format_alt(alt2),
                "contract_number": "",
                "contract_type": "",
                "start_date": "",
                "end_date": "",
                "notes": "AUTO-MATCHED - please verify" if best else "NO MATCH - fill in manually",
            }
            writer.writerow(row)

    stats["register_created"] = True
    print("")
    print(f"Register with suggestions created: {output_path}")
    print("")
    print("Summary:")
    print(f"  Matched:   {stats['matched']} contracts")
    print(f"  Unmatched: {stats['unmatched']} contracts")
    print("")
    print("Next steps:")
    print("  1. Open the CSV and VERIFY suggested matches")
    print("  2. For unmatched contracts, fill in client_id manually")
    print("  3. Remove match_score, match_type, alternative columns before sync")
    print("  4. Run: python sync_contracts.py <register.csv>")

    return stats


def extract_docx(filepath: Path) -> str:
    """
    Extract text content from a Word document.

    Extracts:
    - All paragraphs (preserving structure)
    - Tables (formatted with pipe separators)
    - Headers marked with markdown-style formatting
    """
    doc = Document(str(filepath))
    parts = []

    # Add metadata header
    parts.append(f"# Contract: {filepath.name}")
    parts.append(f"# Extracted: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    parts.append(f"# Source type: Word Document (.docx)")
    parts.append("")
    parts.append("---")
    parts.append("")

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Check if it looks like a heading (short, possibly bold/styled)
            if para.style and "heading" in para.style.name.lower():
                parts.append(f"## {text}")
            else:
                parts.append(text)
            parts.append("")  # Add spacing between paragraphs

    # Extract tables
    if doc.tables:
        parts.append("")
        parts.append("---")
        parts.append("## Tabellen")
        parts.append("")

        for i, table in enumerate(doc.tables, 1):
            parts.append(f"### Tabel {i}")
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.replace("|", "").strip():
                    parts.append(f"| {row_text} |")
            parts.append("")

    return "\n".join(parts)


def extract_xlsx(filepath: Path) -> str:
    """
    Extract text content from an Excel file (all sheets combined).

    Extracts all sheets with clear sheet headers.
    """
    wb = openpyxl.load_workbook(str(filepath), read_only=True, data_only=True)
    parts = []

    # Add metadata header
    parts.append(f"# Contract: {filepath.name}")
    parts.append(f"# Extracted: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    parts.append(f"# Source type: Excel Spreadsheet (.xlsx)")
    parts.append(f"# Sheets: {', '.join(wb.sheetnames)}")
    parts.append("")
    parts.append("---")
    parts.append("")

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        parts.append(f"## Sheet: {sheet_name}")
        parts.append("")

        for row in sheet.iter_rows(values_only=True):
            row_values = [
                str(cell) if cell is not None else ""
                for cell in row
            ]
            if any(v.strip() for v in row_values):
                parts.append("| " + " | ".join(row_values) + " |")

        parts.append("")

    wb.close()
    return "\n".join(parts)


def extract_xlsx_per_sheet(filepath: Path) -> dict[str, str]:
    """
    Extract text content from an Excel file, one output per sheet.

    Returns dict mapping sheet_name -> extracted content.
    """
    wb = openpyxl.load_workbook(str(filepath), read_only=True, data_only=True)
    results = {}

    for sheet_name in wb.sheetnames:
        parts = []

        # Add metadata header
        parts.append(f"# Contract: {sheet_name}")
        parts.append(f"# Source: {filepath.name} (sheet: {sheet_name})")
        parts.append(f"# Extracted: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        parts.append("")
        parts.append("---")
        parts.append("")

        sheet = wb[sheet_name]
        for row in sheet.iter_rows(values_only=True):
            row_values = [
                str(cell) if cell is not None else ""
                for cell in row
            ]
            if any(v.strip() for v in row_values):
                parts.append("| " + " | ".join(row_values) + " |")

        parts.append("")
        results[sheet_name] = "\n".join(parts)

    wb.close()
    return results


def extract_contract(filepath: Path) -> str:
    """Extract contract content based on file type."""
    ext = filepath.suffix.lower()

    if ext == ".docx":
        return extract_docx(filepath)
    elif ext == ".xlsx":
        return extract_xlsx(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .docx, .xlsx")


def get_output_path(source_path: Path, output_folder: Path) -> Path:
    """Generate output path for extracted text file."""
    return output_folder / f"{source_path.stem}.txt"


def extract_contracts(
    source_folder: Path,
    output_folder: Path,
    specific_files: list[str] = None,
    dry_run: bool = False,
    split_sheets: bool = True
) -> dict:
    """
    Extract contracts to text files.

    Args:
        source_folder: Folder containing contract files
        output_folder: Folder to write extracted text files
        specific_files: If provided, only extract these files
        dry_run: If True, show what would be done without writing
        split_sheets: If True, Excel files with multiple sheets become multiple .txt files

    Returns:
        Statistics dict with counts
    """
    stats = {"extracted": 0, "skipped": 0, "errors": 0}

    # Ensure output folder exists
    if not dry_run:
        output_folder.mkdir(parents=True, exist_ok=True)

    # Get files to process
    if specific_files:
        files = [source_folder / f for f in specific_files]
    else:
        files = list(source_folder.glob("*.docx")) + list(source_folder.glob("*.xlsx"))

    if not files:
        print(f"No contract files found in: {source_folder}")
        return stats

    print(f"Found {len(files)} contract file(s) to process")
    print(f"Source folder: {source_folder}")
    print(f"Output folder: {output_folder}")
    print(f"Split Excel sheets: {split_sheets}")
    print("")

    for filepath in sorted(files):
        if not filepath.exists():
            print(f"  ERROR: File not found: {filepath.name}")
            stats["errors"] += 1
            continue

        try:
            ext = filepath.suffix.lower()

            # Handle Excel with sheet splitting
            if ext == ".xlsx" and split_sheets:
                if dry_run:
                    print(f"  [DRY RUN] Would extract sheets from: {filepath.name}")
                    stats["extracted"] += 1
                else:
                    print(f"  Extracting sheets from: {filepath.name}")
                    sheet_contents = extract_xlsx_per_sheet(filepath)

                    for sheet_name, content in sheet_contents.items():
                        # Clean sheet name for filename
                        safe_name = re.sub(r'[^\w\s-]', '', sheet_name).strip()
                        safe_name = re.sub(r'\s+', '_', safe_name)
                        output_path = output_folder / f"{safe_name}.txt"

                        with open(output_path, "w", encoding="utf-8") as f:
                            f.write(content)

                        print(f"    -> {output_path.name} ({len(content)} chars)")
                        stats["extracted"] += 1
            else:
                # Standard extraction (docx or xlsx without splitting)
                output_path = get_output_path(filepath, output_folder)

                if dry_run:
                    print(f"  [DRY RUN] Would extract: {filepath.name} -> {output_path.name}")
                    stats["extracted"] += 1
                else:
                    print(f"  Extracting: {filepath.name}")
                    content = extract_contract(filepath)

                    with open(output_path, "w", encoding="utf-8") as f:
                        f.write(content)

                    print(f"    -> {output_path.name} ({len(content)} chars)")
                    stats["extracted"] += 1

        except Exception as e:
            print(f"  ERROR extracting {filepath.name}: {e}")
            stats["errors"] += 1

    return stats


def main():
    parser = argparse.ArgumentParser(
        description="Extract contracts to LLM-readable text format and generate register"
    )
    parser.add_argument(
        "--source",
        default=None,
        help=f"Source folder with contract files (default: {config.CONTRACTS_FOLDER})"
    )
    parser.add_argument(
        "--output",
        default=None,
        help="Output folder for extracted text files (default: <source>/extracted)"
    )
    parser.add_argument(
        "--files",
        nargs="+",
        help="Specific file(s) to extract (filenames only, not full paths)"
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Show what would be done without writing files"
    )
    parser.add_argument(
        "--register-only",
        action="store_true",
        help="Only generate register template, skip text extraction"
    )
    parser.add_argument(
        "--extract-only",
        action="store_true",
        help="Only extract text files, skip register generation"
    )
    parser.add_argument(
        "--auto-match",
        action="store_true",
        help="Auto-match contracts to Sintes relaties using fuzzy matching"
    )
    parser.add_argument(
        "--min-score",
        type=float,
        default=0.4,
        help="Minimum similarity score for auto-matching (0.0-1.0, default: 0.4)"
    )
    parser.add_argument(
        "--no-split-sheets",
        action="store_true",
        help="Don't split Excel sheets into separate files (default: split enabled)"
    )

    args = parser.parse_args()

    # Set source folder
    source_folder = Path(args.source) if args.source else Path(config.CONTRACTS_FOLDER)
    if not source_folder.exists():
        print(f"ERROR: Source folder not found: {source_folder}")
        sys.exit(1)

    # Set output folder
    if args.output:
        output_folder = Path(args.output)
    else:
        output_folder = source_folder / "llm_conversie"

    print("=" * 60)
    print("Contract Extraction Tool")
    print("=" * 60)
    print(f"Source: {source_folder}")
    print("")

    # Step 1: Generate register template (with or without auto-matching)
    if not args.extract_only:
        print("-" * 60)
        if args.auto_match:
            print("STEP 1: Contract Register with Auto-Matching")
            print("-" * 60)
            register_path = source_folder / "contract_register_matched.csv"
            register_stats = generate_register_with_matching(
                source_folder=source_folder,
                output_path=register_path,
                specific_files=args.files,
                dry_run=args.dry_run,
                min_score=args.min_score
            )
        else:
            print("STEP 1: Contract Register Template")
            print("-" * 60)
            register_path = source_folder / "contract_register_template.csv"
            register_stats = generate_register_template(
                source_folder=source_folder,
                output_path=register_path,
                specific_files=args.files,
                dry_run=args.dry_run
            )
        print("")

    # Step 2: Extract contract text
    if not args.register_only:
        print("-" * 60)
        print("STEP 2: Text Extraction (LLM-readable format)")
        print("-" * 60)
        extract_stats = extract_contracts(
            source_folder=source_folder,
            output_folder=output_folder,
            specific_files=args.files,
            dry_run=args.dry_run,
            split_sheets=not args.no_split_sheets
        )
        print("")

    # Summary
    print("=" * 60)
    print("Summary")
    print("=" * 60)

    if not args.extract_only:
        print(f"Register: {register_stats['files_found']} contracts found")
        if args.auto_match:
            print(f"  Relaties loaded: {register_stats.get('relaties_loaded', 0)}")
            print(f"  Auto-matched: {register_stats.get('matched', 0)}")
            print(f"  Unmatched: {register_stats.get('unmatched', 0)}")
        if register_stats.get("register_created"):
            filename = "contract_register_matched.csv" if args.auto_match else "contract_register_template.csv"
            print(f"  -> Created: {filename}")

    if not args.register_only:
        print(f"Extraction: {extract_stats['extracted']} files extracted, {extract_stats['errors']} errors")
        if extract_stats['extracted'] > 0:
            print(f"  -> Output: {output_folder}")

    if args.dry_run:
        print("")
        print("[DRY RUN] No files were written.")

    print("")
    print("Next steps:")
    print("  1. Fill in client_id in contract_register_template.csv")
    print("  2. Review extracted .txt files in the 'extracted' folder")
    print("  3. Sync register to database: python sync_contracts.py contract_register_template.csv")


if __name__ == "__main__":
    main()
