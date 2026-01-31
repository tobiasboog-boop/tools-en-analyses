"""
Notifica Offerte Generator
===========================

Herbruikbare tool voor het genereren van offertes op basis van Word templates.
Behoudt alle opmaak, logo's en huisstijl.

Gebruik:
    python offerte_generator.py

Auteur: Notifica
Datum: Januari 2026
"""

from docx import Document
from docx.shared import Pt
import re
from datetime import datetime
from pathlib import Path


class OfferteGenerator:
    """Generator voor Notifica offertes op basis van Word templates."""

    def __init__(self, template_path: str):
        """
        Initialiseer de generator met een template.

        Args:
            template_path: Pad naar het Word template bestand
        """
        self.template_path = Path(template_path)
        self.doc = Document(template_path)

    def replace_in_paragraph(self, paragraph, old_text: str, new_text: str):
        """
        Vervang tekst in een paragraph, ook als die over meerdere runs verspreid is.
        """
        if old_text not in paragraph.text:
            return False

        # Probeer eerst simpele per-run vervanging
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return True

        # Als dat niet werkt: combineer runs en vervang
        if paragraph.runs:
            full_text = paragraph.text
            new_full_text = full_text.replace(old_text, new_text)

            # Bewaar opmaak van eerste run, wis de rest
            first_run = paragraph.runs[0]
            for run in paragraph.runs[1:]:
                run.text = ""
            first_run.text = new_full_text
            return True

        return False

    def replace_all(self, replacements: dict):
        """
        Vervang alle placeholders in het document.

        Args:
            replacements: Dict met {oude_tekst: nieuwe_tekst}
        """
        for old_text, new_text in replacements.items():
            # Paragraphs
            for para in self.doc.paragraphs:
                self.replace_in_paragraph(para, old_text, new_text)

            # Tables
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self.replace_in_paragraph(para, old_text, new_text)

            # Headers & Footers
            for section in self.doc.sections:
                for para in section.header.paragraphs:
                    self.replace_in_paragraph(para, old_text, new_text)
                for para in section.footer.paragraphs:
                    self.replace_in_paragraph(para, old_text, new_text)

    def insert_after_paragraph(self, search_text: str, new_paragraphs: list):
        """
        Voeg nieuwe paragrafen toe NA een paragraaf die bepaalde tekst bevat.

        Args:
            search_text: Tekst om te zoeken in paragrafen
            new_paragraphs: Lijst van strings om toe te voegen
        """
        for i, para in enumerate(self.doc.paragraphs):
            if search_text in para.text:
                # Voeg nieuwe paragrafen toe na deze positie
                for j, new_text in enumerate(new_paragraphs):
                    new_para = self.doc.paragraphs[i]._element
                    new_p = new_para.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p', {})
                    new_para.addnext(new_p)
                    # Voeg run toe met tekst
                    from docx.oxml.ns import qn
                    run = new_p.makeelement(qn('w:r'), {})
                    text = run.makeelement(qn('w:t'), {})
                    text.text = new_text
                    run.append(text)
                    new_p.append(run)
                return True
        return False

    def add_investment_section(self, fase1_bedrag="13.800", fase1_dagen="12",
                                opleverdatum="1 maart 2026"):
        """
        Voeg investeringssectie toe aan het document met opgemaakte tabellen.
        """
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        def set_cell_shading(cell, color):
            """Zet achtergrondkleur van een cel."""
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), color)
            cell._tc.get_or_add_tcPr().append(shading)

        def format_header_row(row, bg_color="1F4E79"):
            """Maak header row op met achtergrondkleur en witte tekst."""
            for cell in row.cells:
                set_cell_shading(cell, bg_color)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Pt(10)

        def format_total_row(row, bg_color="D9E2F3"):
            """Maak totaal row op."""
            for cell in row.cells:
                set_cell_shading(cell, bg_color)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)

        # Zoek de juiste paragraaf
        target_idx = None
        for i, para in enumerate(self.doc.paragraphs):
            if "Gefaseerde implementatie" in para.text or "Basis op orde" in para.text:
                target_idx = i
                # Verwijder de oude tekst als die er al staat
                if "INVESTERING" in para.text:
                    para.text = para.text.split("INVESTERING")[0].strip()
                break

        if target_idx is None:
            print("[WAARSCHUWING] Kon geen geschikte plek vinden voor investeringssectie")
            return False

        # Voeg kopje toe
        heading = self.doc.add_paragraph()
        heading_run = heading.add_run(f"\nInvestering Fase 1: Basis op Orde")
        heading_run.bold = True
        heading_run.font.size = Pt(12)

        intro = self.doc.add_paragraph()
        intro_run = intro.add_run(f"Oplevering: {opleverdatum}")
        intro_run.font.size = Pt(10)

        # === TABEL FASE 1 ===
        fase1_data = [
            ("Onderdeel", "Dagen", "Bedrag"),
            ("AccountView API-koppeling", "3", "€ 3.450"),
            ("Financieel inzicht (dashboards)", "2", "€ 2.300"),
            ("Projecten & onderhanden werk", "2", "€ 2.300"),
            ("Uren & productiviteit", "2", "€ 2.300"),
            ("Power Apps (normering)", "2", "€ 2.300"),
            ("Testen & oplevering", "1", "€ 1.150"),
            ("Totaal Fase 1", "12", f"€ {fase1_bedrag}"),
        ]

        table1 = self.doc.add_table(rows=len(fase1_data), cols=3)
        table1.style = 'Table Grid'

        for row_idx, (col1, col2, col3) in enumerate(fase1_data):
            row = table1.rows[row_idx]
            row.cells[0].text = col1
            row.cells[1].text = col2
            row.cells[2].text = col3
            # Uitlijning
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Header opmaak
        format_header_row(table1.rows[0])
        # Totaal opmaak
        format_total_row(table1.rows[-1])

        # Kolom breedtes
        for row in table1.rows:
            row.cells[0].width = Cm(7)
            row.cells[1].width = Cm(2)
            row.cells[2].width = Cm(3)

        # === VERVOLGFASES ===
        heading2 = self.doc.add_paragraph()
        heading2_run = heading2.add_run("\nVervolgfases (indicatief)")
        heading2_run.bold = True
        heading2_run.font.size = Pt(11)

        vervolg_data = [
            ("Fase", "Omschrijving", "Investering"),
            ("Fase 2", "Service & Onderhoud", "€ 4.600 - € 5.750"),
            ("Fase 3", "Commercieel & Sales", "€ 2.300 - € 4.600"),
            ("Fase 4", "Voorraad & Inkoop", "€ 2.300 - € 4.600"),
            ("Totaal", "Alle fases", "€ 23.000 - € 28.750"),
        ]

        table2 = self.doc.add_table(rows=len(vervolg_data), cols=3)
        table2.style = 'Table Grid'

        for row_idx, (col1, col2, col3) in enumerate(vervolg_data):
            row = table2.rows[row_idx]
            row.cells[0].text = col1
            row.cells[1].text = col2
            row.cells[2].text = col3
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        format_header_row(table2.rows[0])
        format_total_row(table2.rows[-1])

        for row in table2.rows:
            row.cells[0].width = Cm(2)
            row.cells[1].width = Cm(5)
            row.cells[2].width = Cm(5)

        # Lege regel na tabellen
        self.doc.add_paragraph()

        print(f"[OK] Investeringssectie met tabellen toegevoegd")
        return True

    def save(self, output_path: str):
        """Sla het aangepaste document op."""
        self.doc.save(output_path)
        print(f"[OK] Offerte opgeslagen: {output_path}")


def genereer_offerte_bras():
    """
    Genereer de offerte voor Bras Elektrotechniek.
    Implementatiemodel met 4 fases en investeringskosten.
    """

    # Paden
    # Gebruik bestaande Bras offerte als basis (gebruiker heeft al aanpassingen gedaan)
    template = r"C:\Users\tobia\OneDrive - Notifica B.V\Documenten - Sharepoint Notifica intern\104. Verkoop\Offerte documenten ENK\Bras Elektrotechniek - Offerte BI Implementatie februari 2026.docx"
    output = r"C:\Users\tobia\OneDrive - Notifica B.V\Documenten - Sharepoint Notifica intern\104. Verkoop\Offerte documenten ENK\Bras Elektrotechniek - Offerte BI Implementatie februari 2026 - v2.docx"

    vandaag = datetime.now().strftime("%d-%m-%Y")

    # Vervangingen - alleen investeringsmodel tekst (klantgegevens zijn al correct)
    replacements = {
        # === AANHEF CORRECTIE ===
        "Geachte heer Bras, beste Peter,": "Geachte heer Bras,",

        # === INVESTERINGSMODEL (RaaS -> Implementatie) ===
        # Planning aanpassingen
        "in week 8": "per 1 maart 2026",
        "week 8": "1 maart 2026",
        "Binnen 4 weken": "Binnen 4 weken na opdrachtverstrekking",

        # Verwijder "onze investering" tekst - klant investeert nu
        "Basis op orde is onze investering": "Gefaseerde implementatie met investering",

        # Verander de zin over geen hoge instapkosten
        "dragen de volledige investering voor de ontsluiting van de data": "leveren een gefaseerde implementatie inclusief ontsluiting van de data",
        "De oplossing wordt geleverd tegen een vast maandtarief, zonder hoge instapkosten.": "De investering voor Fase 1 (Basis op Orde) bedraagt EUR 13.800,- excl. BTW (12 dagen). Daarnaast geldt een maandelijks beheertarief.",

        # Samenvattend sectie aanpassen
        "Notifica investeert substantieel in deze oplossing": "Deze implementatie omvat een complete BI-oplossing in vier fases",
        "en vraagt in ruil slechts een gereduceerde vergoeding": ". Fase 1 wordt opgeleverd per 1 maart 2026",
        "Zo realiseren we samen een snelle, betaalbare en schaalbare BI-oplossing": "Vervolgfases (Service & Onderhoud, Commercieel, Voorraad & Inkoop) worden in overleg gepland. Zo realiseren we samen een complete BI-oplossing",

        # Consultancy dagen sectie
        "twee betaalde consultancy-dagen": "de benodigde consultancy-dagen (onderdeel van de investering)",

        # 50% korting verwijderen (geldt niet bij implementatiemodel)
        "Notifica neemt in het eerste jaar 50% van de benodigde consultancyuren voor haar rekening, ontwikkeling uiteraard altijd in afstemming met en in opdracht van": "Vervolgontwikkeling vindt plaats op basis van consultancy, altijd in afstemming met",
    }

    gen = OfferteGenerator(template)
    gen.replace_all(replacements)
    gen.add_investment_section(
        fase1_bedrag="13.800",
        fase1_dagen="12",
        opleverdatum="1 maart 2026"
    )
    gen.save(output)

    return output


def genereer_offerte(
    template_path: str,
    output_path: str,
    bedrijfsnaam: str,
    contactpersoon: str,
    adres: str,
    postcode_plaats: str,
    aanhef: str,
    boekhoudpakket: str = "AccountView",
    maandtarief: str = "650",
    datum: str = None
):
    """
    Generieke functie om een offerte te genereren.

    Args:
        template_path: Pad naar template
        output_path: Pad voor output
        bedrijfsnaam: Naam van het bedrijf
        contactpersoon: Naam contactpersoon (bijv. "t.a.v. Ing. P. Bras, directeur")
        adres: Straat en huisnummer
        postcode_plaats: Postcode en plaats
        aanhef: Aanhef in de brief (bijv. "Geachte heer Bras,")
        boekhoudpakket: AFAS, AccountView, Twinfield, etc.
        maandtarief: Maandelijks tarief
        datum: Datum (default: vandaag)
    """

    if datum is None:
        datum = datetime.now().strftime("%d-%m-%Y")

    # Standaard template waarden (FCG als basis)
    replacements = {
        "FCG Holding B.V. die": f"{bedrijfsnaam} die",
        "van FCG Holding B.V.": f"van {bedrijfsnaam}",
        "t.a.v. Jean-Paul Jansen": contactpersoon,
        "6657KA  Boven-Leeuwen": postcode_plaats,
        "ENK i.c.m. AFAS": f"ENK i.c.m. {boekhoudpakket}",
        "i.c.m. AFAS": f"i.c.m. {boekhoudpakket}",
        "AFAS-database": f"{boekhoudpakket}-database",
        "Beste Jean-Paul,": aanhef,
        "Noord Zuidweg 10": adres,
        "FCG Holding B.V.": bedrijfsnaam,
        "13-6-2025": datum,
        "AFAS": boekhoudpakket,
        "EUR 650,-": f"EUR {maandtarief},-",
    }

    gen = OfferteGenerator(template_path)
    gen.replace_all(replacements)
    gen.save(output_path)

    return output_path


VOORBEELD_KLANTEN = {
    "bras": {
        "bedrijfsnaam": "BRAS ELEKTROTECHNIEK B.V.",
        "contactpersoon": "t.a.v. Ing. P. Bras, directeur",
        "adres": "Zeeheld 3",
        "postcode_plaats": "5342 VX  OSS",
        "aanhef": "Geachte heer Bras,",
        "boekhoudpakket": "AccountView",
        "maandtarief": "650"
    },
    "teo": {
        "bedrijfsnaam": "TEO ELEKTROTECHNIEK B.V.",
        "contactpersoon": "t.a.v. de directie",
        "adres": "Zeeheld 3",
        "postcode_plaats": "5342 VX  OSS",
        "aanhef": "Geachte heer/mevrouw,",
        "boekhoudpakket": "AccountView",
        "maandtarief": "650"
    }
}


if __name__ == "__main__":
    print("=" * 60)
    print("Notifica Offerte Generator")
    print("=" * 60)

    output = genereer_offerte_bras()

    print(f"\n[OK] Offerte gegenereerd!")
    print(f"   Output: {output}")
